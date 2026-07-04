#!/usr/bin/env python3
"""Regenerate the QA-03 fixture set (contract: tests/fixtures/README.md).

Fixtures are git-ignored (binaries); this script is the source of truth.
Run:  uv run --with reportlab --with python-docx python tests/fixtures/generate.py

Each file exercises one parser/validation branch — see the manifest table.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

HERE = Path(__file__).parent


def write_ok_pdf(path: Path) -> None:
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path))
    c.drawString(72, 720, "Knowledge base sample document.")
    c.drawString(72, 700, "Elasticsearch indexes this passage for search relevance.")
    c.showPage()
    c.save()


def write_fonts_pdf(path: Path) -> None:
    # Non-standard (non-Base14) embedded TrueType font — extraction edge case.
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    # A font guaranteed present on macOS; embedding a non-core font is the point.
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    if not Path(font_path).exists():
        # fallback: still a valid PDF, just without the exotic-font wrinkle
        write_ok_pdf(path)
        return
    pdfmetrics.registerFont(TTFont("ExoticArial", font_path))
    c = canvas.Canvas(str(path))
    c.setFont("ExoticArial", 14)
    c.drawString(72, 720, "Ünïcödé passage with an embedded non-core font — ligature ﬁ.")
    c.showPage()
    c.save()


def write_ok_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Sample DOCX", level=1)
    doc.add_paragraph("This paragraph is the searchable content of the DOCX fixture.")
    doc.add_paragraph("A second chunk to exercise chunk boundaries.")
    doc.save(str(path))


def write_empty_docx(path: Path) -> None:
    # Structurally valid but content-free .docx (valid zip, empty body).
    from docx import Document

    Document().save(str(path))


# QA-05 golden set: one topical document per reference query in
# tests/quality/precision_at_3.py. Text is distinctive enough that BM25 ranks
# each document top-1 for its own query. A few lines also carry the load-test
# HOT_QUERIES terms ("database indexing", "final exam schedule", ...) so the
# QA-04 hot bucket returns non-empty results.
_GOLDEN_PDF: dict[str, str] = {
    "ml-lecture.pdf": (
        "Machine learning lecture notes. This lecture introduces supervised and "
        "unsupervised machine learning, gradient descent, and neural networks. It "
        "also covers database indexing strategies for storing feature vectors."
    ),
    "schedule.pdf": (
        "Final exam schedule for the spring semester. This schedule lists exam "
        "dates, times, and rooms for every course. Check the final exam schedule "
        "before the spring examination week begins."
    ),
    "thesis-guide.pdf": (
        "Thesis submission guidelines. These guidelines describe the thesis "
        "submission process, formatting requirements, and deadlines for graduate "
        "students submitting a thesis to the committee."
    ),
    "library.pdf": (
        "Library opening hours. The university library opening hours are Monday to "
        "Friday from 8am to 10pm. Weekend library hours are shorter. Check opening "
        "hours during public holidays."
    ),
}
_GOLDEN_DOCX: dict[str, str] = {
    "scholarship.docx": (
        "Scholarship application form. Complete this scholarship application form "
        "to apply for merit and need-based financial aid. Submit the scholarship "
        "application before the deadline."
    ),
    "registration.docx": (
        "Course registration deadline. The course registration deadline for the "
        "upcoming semester is approaching. Register for your courses before the "
        "registration deadline."
    ),
}


def write_golden(here: Path) -> None:
    from reportlab.pdfgen import canvas

    for name, text in _GOLDEN_PDF.items():
        c = canvas.Canvas(str(here / name))
        c.drawString(72, 720, text)
        c.showPage()
        c.save()
    for name, text in _GOLDEN_DOCX.items():
        from docx import Document

        doc = Document()
        doc.add_paragraph(text)
        doc.save(str(here / name))


def main() -> None:
    write_ok_pdf(HERE / "ok.pdf")
    write_fonts_pdf(HERE / "fonts.pdf")
    write_ok_docx(HERE / "ok.docx")
    write_empty_docx(HERE / "empty.docx")
    write_golden(HERE)

    # empty.pdf: zero-byte file — trips the size==0 validation branch.
    (HERE / "empty.pdf").write_bytes(b"")

    # corrupt.pdf: PDF magic then garbage — parser must fail gracefully.
    (HERE / "corrupt.pdf").write_bytes(b"%PDF-1.4\n\xde\xad\xbe\xef broken xref\n")

    # corrupt.docx: ZIP magic then garbage — not a valid docx (broken formatting).
    (HERE / "corrupt.docx").write_bytes(b"PK\x03\x04\xde\xad\xbe\xef not a real zip\n")

    # note.txt: unsupported content-type branch.
    (HERE / "note.txt").write_text("Plain text is not an accepted upload type.\n")

    # big.bin: > 20 MB size-limit branch. Sparse-ish zeros, cheap to make.
    with open(HERE / "big.bin", "wb") as f:
        f.seek(21 * 1024 * 1024 - 1)
        f.write(b"\0")

    # Sanity: empty.docx must be a real zip; corrupt.pdf must NOT parse as zip.
    assert zipfile.is_zipfile(HERE / "empty.docx"), "empty.docx is not a valid zip"
    assert not zipfile.is_zipfile(HERE / "corrupt.docx"), "corrupt.docx must be a broken zip"
    assert (HERE / "empty.pdf").stat().st_size == 0
    assert (HERE / "big.bin").stat().st_size > 20 * 1024 * 1024
    print("fixtures written to", HERE)


if __name__ == "__main__":
    main()
